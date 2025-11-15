"""
File parsing utilities for the Multi-Agent Resume Optimizer.

This module provides functionality to parse resume files in various formats
(JSON, PDF, DOCX) and convert them to a standardized JSON format.
"""

import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional, Union
import tempfile
import os

try:
    import PyPDF2
    import pypdf
    from docx import Document
except ImportError as e:
    logging.warning(f"Some file parsing dependencies not available: {e}")

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Unified parser for resume files in multiple formats.
    
    Supports JSON, PDF, and DOCX file formats and converts them to a
    standardized JSON structure for the resume optimizer workflow.
    """
    
    def __init__(self):
        """Initialize the resume parser."""
        self.supported_formats = ['.json', '.pdf', '.docx']
    
    def parse_file(self, file_path: Union[str, Path], file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """
        Parse a resume file and return standardized JSON data.
        
        Args:
            file_path: Path to the file or file name for type detection
            file_content: Optional file content as bytes (for Streamlit uploads)
            
        Returns:
            Dict containing standardized resume data
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist and no content provided
            Exception: If parsing fails
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported: {self.supported_formats}")
        
        if file_content is not None:
            # Handle file content from uploads
            return self._parse_from_content(file_content, file_extension)
        else:
            # Handle file path
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            return self._parse_from_path(file_path)
    
    def _parse_from_content(self, content: bytes, file_extension: str) -> Dict[str, Any]:
        """Parse file content from bytes."""
        if file_extension == '.json':
            return self._parse_json_content(content)
        elif file_extension == '.pdf':
            return self._parse_pdf_content(content)
        elif file_extension == '.docx':
            return self._parse_docx_content(content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _parse_from_path(self, file_path: Path) -> Dict[str, Any]:
        """Parse file from file path."""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.json':
            return self._parse_json_file(file_path)
        elif file_extension == '.pdf':
            return self._parse_pdf_file(file_path)
        elif file_extension == '.docx':
            return self._parse_docx_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _parse_json_content(self, content: bytes) -> Dict[str, Any]:
        """Parse JSON content from bytes."""
        try:
            text_content = content.decode('utf-8')
            return json.loads(text_content)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
        except UnicodeDecodeError as e:
            raise ValueError(f"Invalid text encoding: {e}")
    
    def _parse_json_file(self, file_path: Path) -> Dict[str, Any]:
        """Parse JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format in {file_path}: {e}")
        except UnicodeDecodeError as e:
            raise ValueError(f"Invalid text encoding in {file_path}: {e}")
    
    def _parse_pdf_content(self, content: bytes) -> Dict[str, Any]:
        """Parse PDF content from bytes."""
        try:
            # Try pypdf first (more modern)
            try:
                import io
                pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                return self._extract_resume_data_from_text(text_content)
            except ImportError:
                # Fallback to PyPDF2
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                return self._extract_resume_data_from_text(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")
    
    def _parse_pdf_file(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF file."""
        try:
            # Try pypdf first (more modern)
            try:
                pdf_reader = pypdf.PdfReader(str(file_path))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                return self._extract_resume_data_from_text(text_content)
            except ImportError:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_content = ""
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
                    return self._extract_resume_data_from_text(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF {file_path}: {e}")
    
    def _parse_docx_content(self, content: bytes) -> Dict[str, Any]:
        """Parse DOCX content from bytes."""
        try:
            import io
            doc = Document(io.BytesIO(content))
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return self._extract_resume_data_from_text(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    def _parse_docx_file(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file."""
        try:
            doc = Document(str(file_path))
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return self._extract_resume_data_from_text(text_content)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX {file_path}: {e}")
    
    def _extract_resume_data_from_text(self, text: str) -> Dict[str, Any]:
        """
        Extract structured resume data from plain text.
        
        This is a basic implementation that attempts to identify common
        resume sections. In a production system, this would use more
        sophisticated NLP techniques.
        """
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Initialize resume structure
        resume_data = {
            "personal_info": {
                "name": "",
                "email": "",
                "phone": "",
                "location": "",
                "linkedin": "",
                "github": ""
            },
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "projects": [],
            "certifications": [],
            "languages": []
        }
        
        current_section = None
        current_item = {}
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Detect section headers
            if any(keyword in line_lower for keyword in ['summary', 'objective', 'profile']):
                current_section = 'summary'
                continue
            elif any(keyword in line_lower for keyword in ['experience', 'work history', 'employment']):
                current_section = 'experience'
                continue
            elif any(keyword in line_lower for keyword in ['education', 'academic']):
                current_section = 'education'
                continue
            elif any(keyword in line_lower for keyword in ['skills', 'technical skills', 'competencies']):
                current_section = 'skills'
                continue
            elif any(keyword in line_lower for keyword in ['projects', 'portfolio']):
                current_section = 'projects'
                continue
            elif any(keyword in line_lower for keyword in ['certifications', 'certificates']):
                current_section = 'certifications'
                continue
            elif any(keyword in line_lower for keyword in ['languages']):
                current_section = 'languages'
                continue
            
            # Extract personal information from first few lines
            if i < 5:
                if '@' in line and not resume_data['personal_info']['email']:
                    resume_data['personal_info']['email'] = line
                elif any(char.isdigit() for char in line) and len(line) > 7 and not resume_data['personal_info']['phone']:
                    resume_data['personal_info']['phone'] = line
                elif 'linkedin.com' in line_lower:
                    resume_data['personal_info']['linkedin'] = line
                elif 'github.com' in line_lower:
                    resume_data['personal_info']['github'] = line
                elif not resume_data['personal_info']['name'] and len(line) < 50:
                    resume_data['personal_info']['name'] = line
            
            # Process content based on current section
            if current_section == 'summary' and line and line.upper() not in ['SUMMARY', 'OBJECTIVE', 'PROFILE']:
                resume_data['summary'] += line + " "
            elif current_section == 'skills' and line:
                # Split skills by common delimiters
                skills = [skill.strip() for skill in line.split(',') if skill.strip()]
                resume_data['skills'].extend(skills)
            elif current_section == 'experience' and line:
                # Basic experience parsing (would need more sophisticated logic)
                if any(keyword in line_lower for keyword in ['company', 'corporation', 'inc', 'ltd']):
                    if current_item:
                        resume_data['experience'].append(current_item)
                    current_item = {'company': line, 'position': '', 'duration': '', 'description': ''}
                elif current_item and not current_item['position']:
                    current_item['position'] = line
                elif current_item:
                    current_item['description'] += line + " "
            elif current_section == 'education' and line:
                if any(keyword in line_lower for keyword in ['university', 'college', 'institute', 'school']):
                    if current_item:
                        resume_data['education'].append(current_item)
                    current_item = {'institution': line, 'degree': '', 'year': '', 'gpa': ''}
                elif current_item and not current_item['degree']:
                    current_item['degree'] = line
            elif current_section == 'projects' and line:
                if current_item:
                    if 'description' not in current_item:
                        current_item['description'] = ''
                    current_item['description'] += line + " "
                else:
                    current_item = {'name': line, 'description': '', 'technologies': []}
            elif current_section == 'certifications' and line:
                resume_data['certifications'].append(line)
            elif current_section == 'languages' and line:
                resume_data['languages'].append(line)
        
        # Add the last item if exists
        if current_item and current_section == 'experience':
            resume_data['experience'].append(current_item)
        elif current_item and current_section == 'education':
            resume_data['education'].append(current_item)
        elif current_item and current_section == 'projects':
            resume_data['projects'].append(current_item)
        
        # Clean up summary
        resume_data['summary'] = resume_data['summary'].strip()
        
        # Remove duplicates from skills
        resume_data['skills'] = list(set(resume_data['skills']))
        
        return resume_data
    
    def validate_resume_data(self, resume_data: Dict[str, Any]) -> bool:
        """
        Validate that the parsed resume data has required fields.
        
        Args:
            resume_data: Parsed resume data
            
        Returns:
            True if valid, False otherwise
        """
        required_fields = ['personal_info', 'experience', 'education', 'skills']
        
        for field in required_fields:
            if field not in resume_data:
                logger.warning(f"Missing required field: {field}")
                return False
        
        # Check if personal_info has at least name or email
        personal_info = resume_data.get('personal_info', {})
        if not personal_info.get('name') and not personal_info.get('email'):
            logger.warning("Missing name and email in personal_info")
            return False
        
        return True


def parse_resume_file(file_path: Union[str, Path], file_content: Optional[bytes] = None) -> Dict[str, Any]:
    """
    Convenience function to parse a resume file.
    
    Args:
        file_path: Path to the file or file name for type detection
        file_content: Optional file content as bytes (for Streamlit uploads)
        
    Returns:
        Dict containing standardized resume data
    """
    parser = ResumeParser()
    return parser.parse_file(file_path, file_content)


if __name__ == "__main__":
    """Test the file parser with sample files."""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python file_parser.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        parser = ResumeParser()
        resume_data = parser.parse_file(file_path)
        
        print("✅ File parsed successfully!")
        print(f"📄 File: {file_path}")
        print(f"👤 Name: {resume_data.get('personal_info', {}).get('name', 'N/A')}")
        print(f"📧 Email: {resume_data.get('personal_info', {}).get('email', 'N/A')}")
        print(f"💼 Experience entries: {len(resume_data.get('experience', []))}")
        print(f"🎓 Education entries: {len(resume_data.get('education', []))}")
        print(f"🛠️ Skills: {len(resume_data.get('skills', []))}")
        
        # Validate the data
        is_valid = parser.validate_resume_data(resume_data)
        print(f"✅ Valid: {is_valid}")
        
    except Exception as e:
        print(f"❌ Error parsing file: {e}")
        sys.exit(1)
